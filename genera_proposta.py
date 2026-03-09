from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ASSISTENTE TECNICO INTELLIGENTE\nPER ARIA COMPRESSA')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Proposta Progettuale')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph()  # spacer

# --- INTRO ---
def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)

add_separator(doc)
doc.add_paragraph()

intro = doc.add_paragraph()
run = intro.add_run('Ciao Germano,')
run.bold = True
run.font.size = Pt(11)

intro2 = doc.add_paragraph(
    'questa e\u2019 un\u2019idea di progetto basata sulle caratteristiche della tua azienda. '
    'Una volta strutturato e validato, potremmo pensare di proporlo a tutte le filiali in Italia. '
    'Ovviamente solo se la cosa ti interessa \u2014 pero\u2019 dai un\u2019occhiata perche\u2019 credo ne valga la pena.'
)

doc.add_paragraph()

# --- SECTION: COS'E' ---
def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
    run.font.name = 'Calibri'
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '4',
        qn('w:color'): '1a568e'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

add_section_title(doc, 'Cos\u2019e\u2019 in pratica')

p = doc.add_paragraph(
    'Un assistente digitale accessibile da web e WhatsApp che:'
)

bullets = [
    ('Risponde ai clienti', ' su prodotti, compatibilita\u2019 e manutenzione \u2014 24 ore su 24'),
    ('Dimensiona impianti', ' partendo dai dati del cliente (settore, utensili, ore di lavoro, pressione)'),
    ('Genera preventivi', ' con la configurazione corretta'),
    ('Passa la mano al tecnico', ' quando il caso e\u2019 troppo complesso'),
]

for bold_part, normal_part in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(normal_part)
    run.font.size = Pt(11)

doc.add_paragraph()
note = doc.add_paragraph()
run = note.add_run('Non inventa nulla: ')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = note.add_run('ragiona sui vostri cataloghi, le vostre schede tecniche e le regole dei vostri tecnici.')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# --- SECTION: COME SI FA ---
add_section_title(doc, 'Come si realizza \u2014 4 step')

table = doc.add_table(rows=5, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header
headers = ['Step', 'Cosa si fa', 'Tempo']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'
    # Background color
    shading = cell._element.makeelement(qn('w:shd'), {
        qn('w:fill'): '1a568e',
        qn('w:val'): 'clear'
    })
    cell._element.get_or_add_tcPr().append(shading)

rows_data = [
    ('1', 'Ci date cataloghi, schede e listini (roba che avete gia\u2019).\nFacciamo 2-3 chiacchierate col vostro tecnico piu\u2019 esperto per catturare le regole pratiche', '1 settimana'),
    ('2', 'Noi costruiamo il sistema e lo colleghiamo ai vostri documenti', '2 settimane'),
    ('3', 'Il vostro tecnico verifica che le risposte siano corrette su casi reali. Correggiamo dove sbaglia', '1 settimana'),
    ('4', 'Si accende su una filiale pilota, si monitora e si affina', '1 settimana'),
]

for row_idx, (step, cosa, tempo) in enumerate(rows_data, 1):
    for col_idx, val in enumerate([step, cosa, tempo]):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if col_idx == 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            shading = cell._element.makeelement(qn('w:shd'), {
                qn('w:fill'): 'f0f5fa',
                qn('w:val'): 'clear'
            })
            cell._element.get_or_add_tcPr().append(shading)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(11)
    row.cells[2].width = Cm(3)

doc.add_paragraph()

total = doc.add_paragraph()
total.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = total.add_run('Dal via al primo utilizzo reale: circa 5 settimane')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

doc.add_paragraph()

# --- SECTION: COSA SERVE ---
add_section_title(doc, 'Cosa serve da parte vostra')

items_serve = [
    'I PDF che avete gia\u2019 (cataloghi, schede tecniche, listino)',
    '10-15 ore del tecnico piu\u2019 esperto, distribuite in 2-3 incontri',
    'Una filiale che faccia da pilota',
]

for item in items_serve:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Nessun software da installare, nessun server da comprare. Funziona da qualsiasi dispositivo.')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# --- SECTION: PERCHE' CONVIENE ---
add_section_title(doc, 'Perche\u2019 conviene')

benefits = [
    ('Il tecnico si libera', ' delle domande ripetitive (il 60-70% del totale)'),
    ('I clienti ricevono risposta immediata', ', anche di notte o nel weekend'),
    ('Un nuovo commerciale diventa operativo', ' in giorni invece che in mesi'),
    ('Nessun concorrente in Italia', ' offre questo servizio oggi'),
]

for bold_part, normal_part in benefits:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(normal_part)
    run.font.size = Pt(11)

doc.add_paragraph()

# --- SECTION: ESPANSIONE ---
add_section_title(doc, 'L\u2019idea piu\u2019 grande')

p = doc.add_paragraph(
    'Una volta che funziona sulla tua filiale, ogni altra filiale in Italia si attiva in 2-3 giorni. '
    'La conoscenza tecnica e\u2019 gia\u2019 dentro il sistema \u2014 si personalizza solo il listino e i contatti locali.'
)

doc.add_paragraph()

# Flow diagram as text
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow.add_run('Filiale pilota  \u2192  Regione  \u2192  Tutta Italia')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

flow2 = doc.add_paragraph()
flow2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow2.add_run('(mese 1-2)          (mese 3-4)       (mese 5-6)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
doc.add_paragraph()

# --- CTA ---
add_separator(doc)
doc.add_paragraph()

cta = doc.add_paragraph()
cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cta.add_run('Prossimo passo')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

cta2 = doc.add_paragraph()
cta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cta2.add_run('Ci sentiamo mezz\u2019ora, ti faccio vedere un esempio concreto\ne decidiamo se partire.')
run.font.size = Pt(12)

# Save
doc.save('/home/user/comprossori-germano/Proposta-Germano.docx')
print('DOCX creato con successo')
